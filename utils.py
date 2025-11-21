import os
import docx
from lxml import etree
from pinecone import Pinecone
import streamlit as st
from langchain_openai import OpenAIEmbeddings
# from langchain.document_loaders import PyPDFLoader
# from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_pinecone import PineconeVectorStore

# Pinecone Connection & Operations
@st.cache_resource
def get_pinecone_client():
    """Cache the Pinecone connection to avoid repeated initialization."""
    return Pinecone(api_key=os.environ.get("PINECONE_API_KEY"))

def fetch_pinecone_namespaces(index_name):
    """Retrieve the list of all existing namespaces from Pinecone."""
    pc = get_pinecone_client()
    try:
        index_stats = pc.describe_index(index_name).stats
        return list(index_stats.namespaces.keys()) if index_stats and index_stats.namespaces else []
    except Exception:
        return []

# GCO Experience Extraction Function (from Extract_GCO_wisdom.py)
def extract_revisions_from_single_doc(file_path, nsmap):
    """Directly read the tracked changes from a single .docx file."""
    doc = docx.Document(file_path)
    extracted_data = []
    for para in doc.paragraphs:
        if '<w:ins' in para._p.xml or '<w:del' in para._p.xml:
            original_text, revised_text = "", ""
            p_tree = etree.fromstring(para._p.xml)
            runs = p_tree.xpath('.//w:r', namespaces=nsmap)
            for run in runs:
                text_nodes = run.xpath('.//w:t', namespaces=nsmap)
                text = text_nodes[0].text if text_nodes and text_nodes[0].text else ""
                if run.xpath('.//w:ins', namespaces=nsmap): revised_text += text
                elif run.xpath('.//w:del', namespaces=nsmap): original_text += text
                else:
                    original_text += text
                    revised_text += text
            if original_text.strip() != revised_text.strip():
                wisdom_chunk = (
                    f"【Review Case – Tracked Changes】\n"
                    f"  - Original Text Before Revision：\n"
                    f"    ---\n"
                    f"    {original_text.strip()}\n"
                    f"    ---\n"
                    f"  - Suggested Text After Revision：\n"
                    f"    ---\n"
                    f"    {revised_text.strip()}\n"
                    f"    ---\n"
                )
                extracted_data.append(wisdom_chunk)
    return extracted_data

def extract_comments_from_docx(file_path):
    """Extract all comments and their associated text from a .docx file."""
    doc = docx.Document(file_path)
    extracted_data = []
    for comment in doc.comments:
        paragraphs = comment.paragraphs
        original_text = "\n".join([p.text for p in paragraphs]).strip()
        comment_text = comment.text.strip()
        if original_text and comment_text:
            wisdom_chunk = (
                f"【Review Case – Legal Expert Comments】\n"
                f"  - Associated Original Text：\n"
                f"    ---\n"
                f"    {original_text.strip()}\n"
                f"    ---\n"
                f"  - Expert Comment：\n"
                f"    ---\n"
                f"    {comment_text}\n"
                f"    ---\n"
            )
            extracted_data.append(wisdom_chunk)
    return extracted_data

# Data Upload (Ingest) Function
def ingest_docs_to_pinecone(docs, index_name, namespace):
    """Upload LangChain document chunks to the specified Pinecone namespace."""
    embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
    PineconeVectorStore.from_documents(
        docs,
        embedding=embeddings,
        index_name=index_name,
        namespace=namespace
    )
